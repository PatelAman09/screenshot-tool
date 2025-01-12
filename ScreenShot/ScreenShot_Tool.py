import os
import datetime
import shutil
from tkinter import filedialog, Tk, Label, Entry, Button, RIDGE
from pynput.keyboard import Key, Listener
from docx import Document
from docx.shared import Inches
import pyautogui

# Global Variables
document = Document()
Shortcut = Key.print_screen  # Screenshot Key
exit_combination = {Key.ctrl_l, Key.esc}  # Exit Combination [Left Ctrl + Esc]
default_doc_name = "testingScreenShots"
img_count = 1
currently_pressed = set()
Path = ""
select_folder = ""


# Keyboard Event Handlers
def on_press(key):
    """Handles key press events."""
    check_key(key)
    if key in exit_combination:
        currently_pressed.add(key)
        if currently_pressed == exit_combination:
            listener.stop()
            exit_program()


def on_release(key):
    """Handles key release events."""
    currently_pressed.discard(key)


# Screenshot Functionality
def check_key(key):
    """Checks if the pressed key is the shortcut key."""
    if key == Shortcut:
        save_screenshot()


def save_screenshot():
    """Captures and saves a screenshot, adds a title, and includes it in the document."""
    global img_count
    screenshot_path = os.path.join(Path, "shots", f"{img_count}.png")
    pyautogui.screenshot().save(screenshot_path)
    add_title()
    add_screenshot_to_doc(screenshot_path)
    print(f"Screenshot saved: {screenshot_path}")
    img_count += 1


# GUI for Adding Titles
def add_title():
    """Displays a GUI for the user to input a title for the screenshot."""
    root = Tk()
    root.title("Add Title")
    
    Label(root, text="Enter The Title:", font="Verdana 12").pack()
    title_entry = Entry(root, width=50, font="Verdana 12")
    title_entry.pack()

    def submit_title():
        user_input = title_entry.get()
        document.add_paragraph(user_input)
        root.destroy()

    Button(
        root, text="Submit", font="Verdana 12", bg="#000048", fg="#fefefe",
        activebackground="#000048", activeforeground="#fefefe", relief=RIDGE, command=submit_title
    ).pack()

    root.mainloop()


def add_screenshot_to_doc(img_path):
    """Adds a screenshot image to the Word document."""
    paragraph = document.add_paragraph()
    run = paragraph.add_run()
    run.add_picture(img_path, width=Inches(5.91), height=Inches(3.54))
    document.add_paragraph()  # Add space after the screenshot


# Document and Directory Management
def save_document(file_name):
    """Saves the Word document with the specified file name."""
    document.save(os.path.join(Path, f"{file_name}.docx"))
    print(f"Document saved at: {os.path.join(Path, f'{file_name}.docx')}")


def create_directories():
    """Creates the necessary directories for storing screenshots."""
    global Path
    timestamp = datetime.datetime.now().strftime("%d.%m.%y_%I.%M.%S.%p")
    Path = os.path.join(select_folder, "TestScreenShot", timestamp)
    os.makedirs(os.path.join(Path, "shots"))
    print(f"Directories created at: {Path}")


def select_folder_dialog():
    """Opens a folder selection dialog for the user."""
    global select_folder
    root = Tk()
    root.withdraw()
    select_folder = filedialog.askdirectory()
    if not select_folder:
        print("No folder selected. Exiting.")
        exit()
    os.chdir(select_folder)
    print(f"Selected folder: {select_folder}")


# Program Initialization and Exit
def print_start_message():
    """Displays the start message."""
    print("<-------------------- Screen Shot Tool -------------------->")
    print("Press 'Print Screen' to take a screenshot.")
    print("Press 'Ctrl + Esc' to exit and save the document.")
    print("----------------------------------------------------------\n")


def exit_program():
    """Handles program exit and saves the document."""
    print("<===================================>")
    file_name = input(f"Enter Document Name (default: {default_doc_name}): ").strip()
    bad_chars = [';', ':', '!', "*", "<", ">", "\"", "/", "|", "?"]
    for char in bad_chars:
        file_name = file_name.replace(char, '')

    file_name = file_name if file_name else default_doc_name
    save_document(file_name)
    shutil.rmtree(os.path.join(Path, "shots"))
    print("Temporary files cleaned up. Exiting.")


# Main Execution
if __name__ == "__main__":
    print_start_message()
    select_folder_dialog()
    create_directories()

    with Listener(on_press=on_press, on_release=on_release) as listener:
        listener.join()
